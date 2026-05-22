"""Load text from PDF and DOCX files under docs/."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path


from docx import Document
from pypdf import PdfReader


@dataclass
class LoadedDocument:
    source_file: str
    pages: list[str]


def load_pdf(path: Path) -> LoadedDocument:
    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        text = (page.extract_text() or "").strip()
        pages.append(text)
    return LoadedDocument(source_file=path.name, pages=pages)


def load_docx(path: Path) -> LoadedDocument:
    doc = Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)
    return LoadedDocument(source_file=path.name, pages=[full_text])


def load_document(path: Path) -> LoadedDocument:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return load_pdf(path)
    if suffix in (".docx", ".doc"):
        return load_docx(path)
    raise ValueError(f"Unsupported file type: {path}")


def _is_valid_doc_path(path: Path) -> bool:
    name = path.name
    if name.startswith(".") or name.startswith("~$"):
        return False
    return path.suffix.lower() in (".pdf", ".docx", ".doc")


def load_all_documents(docs_dir: Path) -> list[LoadedDocument]:
    if not docs_dir.is_dir():
        raise FileNotFoundError(f"Docs directory not found: {docs_dir}")
    documents: list[LoadedDocument] = []
    for path in sorted(docs_dir.iterdir()):
        if not _is_valid_doc_path(path):
            continue
        documents.append(load_document(path))
    return documents


def build_full_corpus(
    docs_dir: Path,
    max_chars: int | None = None,
) -> tuple[str, dict]:
    """Concatenate all documents for full-chart prompt injection."""
    documents = load_all_documents(docs_dir)
    parts: list[str] = []
    for doc in documents:
        body = "\n\n".join(p for p in doc.pages if p.strip())
        parts.append(f"=== DOCUMENT: {doc.source_file} ===\n{body}")

    corpus = "\n\n".join(parts)
    meta = {
        "document_count": len(documents),
        "source_files": [d.source_file for d in documents],
        "total_chars": len(corpus),
        "truncated": False,
        "truncated_chars": 0,
    }

    if max_chars and len(corpus) > max_chars:
        meta["truncated"] = True
        meta["truncated_chars"] = max_chars
        corpus = (
            corpus[:max_chars]
            + "\n\n[... corpus truncated for model context limit ...]"
        )

    return corpus, meta
