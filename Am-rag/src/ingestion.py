from pathlib import Path
from typing import List
import fitz  # pymupdf
import docx
from llama_index.core import Document

from config import DOCS_DIR


def load_pdf(path: Path) -> List[Document]:
    """Load a PDF, one Document per page with metadata."""
    doc = fitz.open(str(path))
    documents = []
    for i, page in enumerate(doc):
        text = page.get_text("text").strip()
        if not text:
            continue
        documents.append(
            Document(
                text=text,
                metadata={
                    "source": path.name,
                    "file_type": "pdf",
                    "page": i + 1,
                    "total_pages": len(doc),
                },
            )
        )
    doc.close()
    return documents


def load_docx(path: Path) -> List[Document]:
    """Load a DOCX as a single Document with section metadata when possible."""
    d = docx.Document(str(path))
    paragraphs = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)
    return [
        Document(
            text=full_text,
            metadata={
                "source": path.name,
                "file_type": "docx",
            },
        )
    ]


def load_all_documents(docs_dir: Path = DOCS_DIR) -> List[Document]:
    """Scan docs_dir and load every supported file."""
    documents = []
    for path in sorted(docs_dir.iterdir()):
        if path.suffix.lower() == ".pdf":
            documents.extend(load_pdf(path))
        elif path.suffix.lower() == ".docx":
            documents.extend(load_docx(path))
        else:
            print(f"Skipped (unsupported): {path.name}")
    return documents


if __name__ == "__main__":
    docs = load_all_documents()
    print(f"Loaded {len(docs)} documents")
    for d in docs[:3]:
        print(f"- {d.metadata['source']} (page {d.metadata.get('page', '-')}) — {len(d.text)} chars")